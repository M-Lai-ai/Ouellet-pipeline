# content_processor.py

import os
import re
import requests
import logging
import hashlib
from datetime import datetime
import time
import warnings

import PyPDF2
import pytesseract
from PIL import Image
from docx import Document
from dotenv import load_dotenv
from pdf2image import convert_from_path
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry
from urllib3.exceptions import NotOpenSSLWarning
import tiktoken  # Assurez-vous d'avoir installé tiktoken

# Filtrer l'avertissement NotOpenSSLWarning (solution temporaire)
warnings.simplefilter('ignore', NotOpenSSLWarning)

# Charger les variables d'environnement
load_dotenv()

# Configuration des clés API
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
if not OPENAI_API_KEY:
    raise EnvironmentError("Veuillez définir la clé API OpenAI dans le fichier .env")

# Configuration du logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("pipeline.log", encoding='utf-8'),
        logging.StreamHandler()
    ]
)

class ContentProcessor:
    def __init__(self, base_dir="crawler_output"):
        self.base_dir = base_dir
        self.create_directories()
        self.session = self.create_requests_session()
        self.processed_files = set()  # Initialiser un set pour les fichiers traités
        self.load_processed_files()

    def create_directories(self):
        """Crée la structure de dossiers nécessaire."""
        directories = ['content', 'logs']
        for dir_name in directories:
            path = os.path.join(self.base_dir, dir_name)
            try:
                os.makedirs(path, exist_ok=True)
                logging.debug(f"Répertoire créé ou déjà existant : {path}")
            except OSError as e:
                logging.error(f"Impossible de créer le répertoire {path}: {e}")
                raise

    def create_requests_session(self):
        """Crée une session requests avec gestion des retries."""
        session = requests.Session()
        retries = Retry(
            total=5,
            backoff_factor=1,
            status_forcelist=[429, 500, 502, 503, 504],
            allowed_methods=["HEAD", "GET", "OPTIONS", "POST"]
        )
        adapter = HTTPAdapter(max_retries=retries)
        session.mount('https://', adapter)
        session.mount('http://', adapter)
        return session

    def sanitize_filename(self, url, file_type, extension, page_number=None):
        """Crée un nom de fichier sécurisé."""
        url_hash = hashlib.md5(url.encode('utf-8')).hexdigest()[:8]
        filename = os.path.basename(url)
        if not filename:
            filename = 'index'
        filename = re.sub(r'[^\w\-_.]', '_', filename)
        name, _ = os.path.splitext(filename)
        if page_number is not None:
            sanitized = f"{name}_page_{page_number:03d}_{url_hash}{extension}"
        else:
            sanitized = f"{name}_{url_hash}{extension}"
        logging.debug(f"Nom de fichier sanitizé: {sanitized}")
        return sanitized

    def extract_text_from_pdf(self, file_path):
        """Extrait le contenu de chaque page d'un PDF."""
        pages_content = []
        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                num_pages = len(reader.pages)
                logging.info(f"Nombre de pages dans {file_path}: {num_pages}")
                for page_number in range(num_pages):
                    page = reader.pages[page_number]
                    text = page.extract_text() or ""
                    
                    # Si le texte est vide, essayer l'OCR
                    if not text.strip():
                        logging.debug(f"Texte vide sur la page {page_number + 1}, tentative d'OCR.")
                        text = self.extract_text_with_ocr(file_path, page_number + 1)
                    
                    if text.strip():
                        # Ajouter des marqueurs pour aider GPT à identifier les structures
                        formatted_text = (
                            f"### Page {page_number + 1}\n\n"
                            f"[DÉBUT CONTENU]\n{text}\n[FIN CONTENU]\n\n"
                            f"Note: Si ce contenu contient des tableaux, "
                            f"veuillez les restructurer en répétant les en-têtes pour chaque entrée."
                        )
                        pages_content.append({
                            'page': page_number + 1,
                            'text': formatted_text
                        })
        except PyPDF2.errors.PdfReadError as e:
            logging.error(f"Erreur de lecture du PDF {file_path}: {e}")
        except Exception as e:
            logging.error(f"Erreur lors de l'extraction du PDF {file_path}: {e}")
        return pages_content

    def extract_text_with_ocr(self, file_path, page_number):
        """Extrait le texte via OCR avec une attention particulière aux tableaux."""
        try:
            images = convert_from_path(file_path, first_page=page_number, last_page=page_number)
            if images:
                # Configuration OCR optimisée pour la détection des tableaux
                custom_config = r'--oem 3 --psm 6'
                text = pytesseract.image_to_string(images[0], lang='fra', config=custom_config)
                logging.debug(f"OCR réussi pour la page {page_number} de {file_path}")
                return text
            logging.warning(f"Aucune image extraite pour la page {page_number} de {file_path}")
            return ""
        except Exception as e:
            logging.warning(f"OCR a échoué pour la page {page_number} de {file_path}: {e}")
            return ""

    def extract_text_from_docx(self, file_path):
        """Extrait le texte des fichiers DOCX avec attention aux tableaux."""
        pages_content = []
        try:
            doc = Document(file_path)
            current_text = []
            page_number = 1  # DOCX n'a pas de pages strictes, donc nous simulons avec un numéro de page unique
            for element in doc.element.body:
                if element.tag.endswith('p'):
                    paragraph = Document().paragraphs[0]
                    paragraph._element = element
                    current_text.append(paragraph.text)
                elif element.tag.endswith('tbl'):
                    table = Document().tables[0]
                    table._element = element
                    current_text.append("\n[DÉBUT TABLEAU]")
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells]
                        current_text.append(" | ".join(row_text))
                    current_text.append("[FIN TABLEAU]\n")
            formatted_text = "\n".join(current_text)
            pages_content.append({
                'page': page_number,
                'text': formatted_text
            })
            logging.debug(f"Extraction DOCX réussie pour {file_path}")
        except Exception as e:
            logging.error(f"Erreur lors de l'extraction du DOCX {file_path}: {e}")
        return pages_content

    def count_tokens(self, text, model="gpt-4"):
        """Compte le nombre de tokens dans le texte."""
        encoding = tiktoken.encoding_for_model(model)
        return len(encoding.encode(text))

    def send_to_gpt(self, content, max_chunk_tokens=8000):
        """Envoie le contenu à GPT en le découpant en blocs basés sur les tokens."""
        headers = {
            'Content-Type': 'application/json',
            'Authorization': f'Bearer {OPENAI_API_KEY}'
        }

        system_prompt = """
        Vous êtes un assistant spécialisé dans la restructuration de documents pour Ouellet Canada.
        Votre tâche est de :

        1. PRÉSERVER ABSOLUMENT TOUT LE CONTENU du document original, sans rien omettre.
        2. Restructurer le contenu en format Markdown clair et lisible.
        3. Pour les tableaux :
           - Identifier les tables correctement délimitées avec [DÉBUT TABLEAU] et [FIN TABLEAU].
           - Transformer chaque ligne en une entrée structurée.
           - Répéter les en-têtes de colonnes pour chaque entrée.
           - Convertir les tables en format Markdown.

        ### Exemple de Transformation :

        **Avant :**
        [DÉBUT TABLEAU]
        Produit Description Quantité Prix
        EUH02B21T 2000W, 208V, 1ph, amande 31 90.00 $
        EUH02B71CT 2000W, 480V, 1ph, amande 18 90.00 $
        [FIN TABLEAU]

        **Après :**
        # Produit: EUH02B21T
        - Description: 2000W, 208V, 1ph, amande
        - Quantité: 31
        - Prix: 90.00 $

        # Produit: EUH02B71CT
        - Description: 2000W, 480V, 1ph, amande
        - Quantité: 18
        - Prix: 90.00 $
        """

        encoding = tiktoken.encoding_for_model("gpt-4")
        tokens = encoding.encode(content)
        chunks = []
        for i in range(0, len(tokens), max_chunk_tokens):
            chunk = encoding.decode(tokens[i:i + max_chunk_tokens])
            chunks.append(chunk)

        restructured_content = ""

        for idx, chunk in enumerate(chunks):
            payload = {
                "model": "gpt-4o-mini",
                "messages": [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": chunk}
                ],
                "temperature": 0,
                "max_tokens": 8000
            }

            # Calculer et loguer la taille du payload
            payload_size = len(str(payload))
            logging.debug(f"Taille du payload envoyé à GPT (bloc {idx + 1}/{len(chunks)}): {payload_size} caractères")

            logging.info(f"Envoi du bloc {idx + 1}/{len(chunks)} à GPT")

            try:
                response = self.session.post(
                    'https://api.openai.com/v1/chat/completions',
                    headers=headers,
                    json=payload,
                    timeout=60
                )
                response.raise_for_status()
                gpt_content = response.json()['choices'][0]['message']['content']
                restructured_content += gpt_content + "\n"
                logging.debug(f"Réponse GPT pour le bloc {idx + 1} reçue avec succès.")
            except requests.exceptions.HTTPError as e:
                # Loguer le statut et le contenu de la réponse
                logging.error(f"Erreur API GPT pour le bloc {idx + 1}: {e} - Statut: {response.status_code} - Réponse: {response.text}")
                continue
            except requests.exceptions.RequestException as e:
                logging.error(f"Erreur API GPT pour le bloc {idx + 1}: {e}")
                continue
            except KeyError:
                logging.error(f"Format de réponse GPT inattendu pour le bloc {idx + 1}.")
                continue

        return restructured_content if restructured_content else None

    def process_files_in_directory(self, directory):
        """Traite tous les fichiers PDF et DOCX dans un répertoire donné."""
        logging.info(f"Processing files in directory: {directory}")
        for root, _, files in os.walk(directory):
            for file in files:
                file_ext = os.path.splitext(file)[1].lower()
                if file_ext in ['.pdf', '.doc', '.docx']:
                    file_path = os.path.join(root, file)
                    logging.info(f"Processing file: {file_path}")
                    try:
                        if file_ext == '.pdf':
                            pages_content = self.extract_text_from_pdf(file_path)
                        elif file_ext in ['.doc', '.docx']:
                            pages_content = self.extract_text_from_docx(file_path)
                        else:
                            logging.warning(f"Unsupported file type: {file_ext} for file {file_path}")
                            continue

                        if not pages_content:
                            logging.warning(f"No content extracted from {file_path}")
                            continue

                        for page_data in pages_content:
                            if not page_data.get('text'):
                                continue

                            # Générer le nom de fichier restructuré
                            filename = self.sanitize_filename(
                                os.path.basename(file_path),
                                'Doc',
                                '_restructured.txt',
                                page_number=page_data['page']  # Inclure le numéro de page
                            )
                            save_path = os.path.join(self.base_dir, 'content', filename)

                            # Vérifier si le fichier restructuré existe déjà
                            if os.path.exists(save_path):
                                logging.info(f"Fichier restructuré déjà existant, skipping: {filename}")
                                continue

                            # Ajouter des marqueurs de contexte pour GPT
                            context = (
                                f"[DÉBUT PAGE {page_data['page']}]\n"
                                f"{page_data['text']}\n"
                                f"[FIN PAGE {page_data['page']}]\n\n"
                            )

                            # Loguer la taille du contenu
                            logging.debug(f"Taille du contenu de la page {page_data['page']} à envoyer: {len(context)} caractères")

                            # Envoyer à GPT pour restructuration
                            restructured_content = self.send_to_gpt(context)
                            if not restructured_content:
                                logging.error(f"Échec de la restructuration pour {file_path} page {page_data['page']}")
                                continue

                            # Sauvegarder le résultat en .txt
                            try:
                                with open(save_path, 'w', encoding='utf-8') as f:
                                    # Ajouter des métadonnées au début du fichier
                                    metadata = (
                                        f"---\n"
                                        f"source_file: {os.path.basename(file_path)}\n"
                                        f"page_number: {page_data['page']}\n"
                                        f"processed_date: {datetime.now().isoformat()}\n"
                                        f"---\n\n"
                                    )
                                    f.write(metadata + restructured_content)
                                logging.info(f"Contenu restructuré sauvegardé dans : {save_path}")

                                # Sauvegarder aussi le contenu brut pour référence en .txt
                                raw_filename = self.sanitize_filename(
                                    os.path.basename(file_path),
                                    'Doc',
                                    '_raw_page.txt',
                                    page_number=page_data['page']
                                )
                                raw_save_path = os.path.join(self.base_dir, 'content', raw_filename)
                                with open(raw_save_path, 'w', encoding='utf-8') as f:
                                    f.write(page_data['text'])
                                logging.info(f"Contenu brut sauvegardé dans : {raw_save_path}")

                                # Sauvegarder le fichier traité dans le fichier de suivi
                                self.save_processed_file(save_path)

                            except IOError as e:
                                logging.error(f"Erreur de sauvegarde pour {file_path} page {page_data['page']}: {e}")

                    except Exception as e:
                        logging.error(f"Erreur lors du traitement de {file_path}: {e}", exc_info=True)

    def load_processed_files(self):
        """Charge les noms des fichiers déjà traités depuis le fichier de suivi."""
        processed_files_path = os.path.join(self.base_dir, 'logs', 'processed_files.txt')
        if os.path.exists(processed_files_path):
            with open(processed_files_path, 'r', encoding='utf-8') as f:
                for line in f:
                    # Stocker les noms de fichiers traités sans le chemin
                    self.processed_files.add(os.path.basename(line.strip()))
            logging.info(f"Loaded {len(self.processed_files)} processed files from tracking file.")
        else:
            logging.info("No processed files tracking file found, starting fresh.")

    def save_processed_file(self, filename):
        """Sauvegarde le nom d'un fichier traité dans le fichier de suivi."""
        processed_files_path = os.path.join(self.base_dir, 'logs', 'processed_files.txt')
        try:
            with open(processed_files_path, 'a', encoding='utf-8') as f:
                f.write(filename + '\n')
            self.processed_files.add(os.path.basename(filename))
            logging.debug(f"Added {filename} to processed files tracking.")
        except Exception as e:
            logging.error(f"Error saving processed file tracking for {filename}: {str(e)}")

    def run_pipeline(self, pdf_directory, doc_directory):
        """Exécute le pipeline de traitement sur les répertoires PDF et Doc."""
        logging.info("Starting ContentProcessor pipeline")
        self.process_files_in_directory(pdf_directory)
        self.process_files_in_directory(doc_directory)
        logging.info("ContentProcessor pipeline completed")
